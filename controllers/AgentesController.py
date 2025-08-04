from flask import render_template, request, redirect, url_for, flash, jsonify, make_response, current_app
from datetime import datetime
from models.models import EquipeInteligente, Usuario, TemplateArquivo, ExecucaoEquipe
from config import db
import json
import base64
import os
import requests
import controllers.IaController as IaController
from PyPDF2 import PdfReader, PdfWriter
import tempfile
import csv
import io
from docx import Document
from docx.shared import Inches
import markdown
import networkx as nx
from io import BytesIO
import pandas as pd
from sqlalchemy import LargeBinary
import matplotlib.pyplot as plt


class AgentesController:
    @staticmethod
    def canva(equipe_id=None):
        """Renderiza a página do canvas para criação ou edição de equipes"""
        equipe_data = None
        if equipe_id:
            equipe = EquipeInteligente.query.get_or_404(equipe_id)
            equipe_data = {
                'id': equipe.id,
                'nome': equipe.nome,
                'descricao': equipe.descricao,
                'processo': equipe.processo,
                'layout': json.loads(equipe.layout) if equipe.layout else {'nodes': [], 'connections': []}
            }

        return render_template('agentes/canva.html', equipe_data=equipe_data)

    @staticmethod
    def criar_equipe_inteligente():
        """Cria uma nova equipe inteligente a partir dos dados do canvas"""
        data = request.get_json()

        try:
            # 1. Criar a equipe
            nova_equipe = EquipeInteligente(
                nome=data.get('nome'),
                descricao=data.get('descricao'),
                processo=data.get('process_type'),
                layout=json.dumps(data.get('layout'))
            )
            db.session.add(nova_equipe)
            db.session.commit()

            return jsonify({'status': 'success', 'message': 'Equipe criada com sucesso!', 'equipe_id': nova_equipe.id})

        except Exception as e:
            db.session.rollback()
            print(f"Erro ao salvar equipe: {e}")
            return jsonify({'status': 'error', 'message': f'Erro ao salvar equipe: {str(e)}'}), 500

    @staticmethod
    def editar_equipe_inteligente(equipe_id):
        """Atualiza uma equipe inteligente existente."""
        equipe = EquipeInteligente.query.get_or_404(equipe_id)
        data = request.get_json()

        try:
            equipe.nome = data.get('nome')
            equipe.descricao = data.get('descricao')
            equipe.processo = data.get('process_type')
            equipe.layout = json.dumps(data.get('layout'))

            db.session.commit()

            return jsonify({'status': 'success', 'message': 'Equipe atualizada com sucesso!', 'equipe_id': equipe.id})

        except Exception as e:
            db.session.rollback()
            print(f"Erro ao atualizar equipe: {e}")
            return jsonify({'status': 'error', 'message': f'Erro ao atualizar equipe: {str(e)}'}), 500

    @staticmethod
    def equipes():
        """Página para listar todas as equipes"""
        equipes = EquipeInteligente.query.all()
        return render_template('agentes/equipes.html', equipes=equipes)

    @staticmethod
    def salvar_menu(equipe_id):
        """Salva a equipe no menu"""
        equipe = EquipeInteligente.query.get(equipe_id)
        if equipe.menu_ordem:
            equipe.menu_ordem = 0
        else:
            equipe.menu_ordem = 1
        db.session.commit()
        return jsonify({'status': 'success', 'message': 'Equipe salva no menu com sucesso!'})

    @staticmethod
    def executar_tarefas(equipe_id, trigger=True):
        equipe = EquipeInteligente.query.get_or_404(equipe_id)
        layout = json.loads(equipe.layout) if isinstance(equipe.layout, str) else equipe.layout

        if trigger:
            return render_template('agentes/executar_tarefas.html', equipe=equipe, trigger=trigger, layout=layout)

        # Tipo de Equipe
        tipo_equipe = layout['nodes'][0]['config']['type']

        # Imprimir o arquivo recebido
        if request.method == 'POST':

            respostas = []

            # Obter o usuário atual
            usuario = Usuario.query.get(1)

            texto = f"""<meta_parametros>
            Meu nome: {usuario.nome}
            Meu email: {usuario.email}
            Data: {datetime.now().strftime('%d/%m/%Y')}
            Hora: {datetime.now().strftime('%H:%M:%S')}
            </meta_parametros>
            
            """

            # Obter o último número de sprint para esta equipe
            ultima_execucao = ExecucaoEquipe.query.filter_by(equipe_id=equipe_id).order_by(ExecucaoEquipe.sprint.desc()).first()
            numero_sprint = (ultima_execucao.sprint or 0) + 1 if ultima_execucao else 1

            # Criar registro de execução no banco de dados
            execucao = ExecucaoEquipe(
                equipe_id=equipe_id,
                sprint=numero_sprint,
                contexto=texto,
                resposta="Execução iniciada"
            )
            db.session.add(execucao)
            db.session.commit()

            ########################################################################################
            # TRIGGER
            ########################################################################################
            if tipo_equipe == 'text':
                texto += request.form['texto']

            else:
                """
                Nós que contém arquivos
                """
                arquivo = request.files['arquivo']

                if tipo_equipe == 'pdf':
                    # Ler o conteúdo do arquivo diretamente do objeto FileStorage
                    pdf_reader = PdfReader(arquivo)

                    for page in pdf_reader.pages:
                        texto += page.extract_text()
                elif tipo_equipe == 'audio':
                    # Salvar o arquivo de áudio temporariamente no disco
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_file:
                        # Ler o conteúdo do arquivo e escrever no arquivo temporário
                        audio_content = arquivo.read()
                        temp_file.write(audio_content)
                        temp_audio_path = temp_file.name  # Caminho do arquivo temporário

                    with open(temp_audio_path, "rb") as file:
                        texto = IaController.groq_transcrever(file)
                    
                    # Limpar arquivo temporário
                    try:
                        os.unlink(temp_audio_path)
                    except:
                        pass
                elif tipo_equipe == 'planilha':
                    # Processar planilha com pandas
                    try:
                        # Obter a extensão do arquivo
                        filename = arquivo.filename.lower()
                        
                        if filename.endswith('.csv'):
                            # Ler arquivo CSV
                            df = pd.read_csv(arquivo)
                        elif filename.endswith(('.xlsx', '.xls')):
                            # Ler arquivo Excel
                            df = pd.read_excel(arquivo)
                        else:
                            raise ValueError("Formato de arquivo não suportado")
                        
                        # Converter DataFrame para texto estruturado
                        texto += f"""
### DADOS DA PLANILHA ###
<informacoes_planilha>
Nome do arquivo: {arquivo.filename}
Dimensões: {df.shape[0]} linhas x {df.shape[1]} colunas
Colunas: {', '.join(df.columns.tolist())}
</informacoes_planilha>

<dados_planilha>
{df.to_string(index=False)}
</dados_planilha>

<resumo_estatistico>
{df.describe().to_string()}
</resumo_estatistico>
"""
                        
                    except Exception as e:
                        texto += f"Erro ao processar planilha: {str(e)}"

            ########################################################################################
            # PROCESSA TODA A TAREFA, EM SEQUENCIA BASEADA NAS CONEXÕES
            ########################################################################################
            
            # Determinar a ordem de execução baseada nas conexões
            def determinar_ordem_execucao(nodes, connections):
                """Determina a ordem correta de execução baseada nas conexões"""
                # Criar um grafo direcionado
                grafo = {}
                nos_entrada = set()
                nos_saida = set()
                
                # Inicializar o grafo
                for node in nodes:
                    grafo[node['id']] = {'dependencias': set(), 'dependentes': set()}
                
                # Processar conexões
                for connection in connections:
                    start_node = connection['startNode']
                    end_node = connection['endNode']
                    
                    # Adicionar dependência
                    grafo[end_node]['dependencias'].add(start_node)
                    grafo[start_node]['dependentes'].add(end_node)
                    
                    nos_entrada.add(start_node)
                    nos_saida.add(end_node)
                
                # Encontrar nós iniciais (que não têm dependências)
                nos_iniciais = set()
                for node_id, info in grafo.items():
                    if not info['dependencias']:
                        nos_iniciais.add(node_id)
                
                # Ordenação topológica
                ordem_execucao = []
                nos_processados = set()
                
                def processar_no(node_id):
                    if node_id in nos_processados:
                        return
                    
                    # Verificar se todas as dependências foram processadas
                    for dep in grafo[node_id]['dependencias']:
                        if dep not in nos_processados:
                            return
                    
                    ordem_execucao.append(node_id)
                    nos_processados.add(node_id)
                    
                    # Processar dependentes
                    for dep in grafo[node_id]['dependentes']:
                        processar_no(dep)
                
                # Processar todos os nós iniciais
                for node_id in nos_iniciais:
                    processar_no(node_id)
                
                # Se ainda há nós não processados, adicionar no final
                for node in nodes:
                    if node['id'] not in ordem_execucao:
                        ordem_execucao.append(node['id'])
                
                return ordem_execucao
            
            # Obter a ordem correta de execução
            ordem_execucao = determinar_ordem_execucao(layout['nodes'], layout['connections'])
            print(f"DEBUG: Ordem de execução: {ordem_execucao}")
            
            # Processar cada nó na ordem determinada pelas conexões
            for node_id in ordem_execucao:
                # Encontrar o nó correspondente
                node = None
                for n in layout['nodes']:
                    if n['id'] == node_id:
                        node = n
                        break
                
                if not node:
                    continue
                
                if node['type'] == 'agent':
                    # Montar o prompt para o agente
                    prompt = f"""{node['config']['backstory']}
                    ### TAREFA ###
                    <tarefa>
                    {node['config']['goal']}
                    </tarefa>

                    ### TEXTO PARA PROCESSAR ###
                    <texto_processar>
                    {texto}
                    </texto_processar>
                    """

                    # Testar primeiro sem o arquivo base64
                    resultado = IaController.groq(prompt)

                    # Extrair a resposta do resultado
                    if isinstance(resultado, tuple) and len(resultado) == 2:
                        if hasattr(resultado[0], 'get_json'):
                            resultado_data = resultado[0].get_json()
                        else:
                            resultado_data = resultado[0]
                    elif isinstance(resultado, dict):
                        resultado_data = resultado
                    elif hasattr(resultado, 'get_json'):
                        resultado_data = resultado.get_json()
                    else:
                        resultado_data = resultado

                    if isinstance(resultado_data, dict) and 'resposta' in resultado_data:
                        resposta_texto = resultado_data['resposta']
                    else:
                        resposta_texto = str(resultado_data)

                    # Criar registro de execução para este agente
                    execucao_agente = ExecucaoEquipe(
                        equipe_id=equipe_id,
                        sprint=execucao.sprint,
                        contexto=prompt,  # Salvar o prompt completo como contexto
                        resposta=resposta_texto
                    )
                    db.session.add(execucao_agente)
                    db.session.commit()

                    print(f"DEBUG - Resposta do agente {node['id']}: {resposta_texto[:200]}...")
                    texto += f'<node_{node["id"]}>{resposta_texto}</node_{node["id"]}>\n\n\n'
                    respostas.append({
                        'resposta': resposta_texto,
                        'node': node
                    })


                ########################################################################################
                # TAREFA
                ########################################################################################
                if node['type'] == 'task':
                    # Obter a última resposta e garantir que seja uma string
                    ultima_resposta = respostas[-1]['resposta']

                    # Se for uma tupla (Response, status_code), extrair o conteúdo
                    if isinstance(ultima_resposta, tuple):
                        if hasattr(ultima_resposta[0], 'get_data'):
                            texto_entrada = ultima_resposta[0].get_data(as_text=True)
                        else:
                            texto_entrada = str(ultima_resposta[0])
                    else:
                        texto_entrada = str(ultima_resposta)

                    # Obter o código Python da configuração da task
                    codigo_python = node['config']['expectedOutput']
                    texto_entrada = '00008323520184013202'  # Valor fixo para teste

                    # Preparar o código completo
                    codigo_completo = codigo_python.replace('{texto_entrada}', texto_entrada)

                    # Criar o contexto da task
                    contexto_task = f"""### TAREFA PYTHON ###
                    <entrada>
                    {texto_entrada}
                    </entrada>

                    <codigo>
                    {codigo_python}
                    </codigo>"""

                    # Executar o código
                    local_vars = {}
                    exec(codigo_completo, globals(), local_vars)
                    resultado_task = local_vars.get('resultado', 'Código executado com sucesso')

                    # Criar registro de execução para esta task
                    execucao_task = ExecucaoEquipe(
                        equipe_id=equipe_id,
                        sprint=execucao.sprint,
                        contexto=contexto_task,
                        resposta=str(resultado_task)
                    )
                    db.session.add(execucao_task)
                    db.session.commit()

                    # Adicionar o resultado ao array de respostas
                    respostas.append({
                        'resposta': resultado_task,
                        'node': node
                    })

                    # Atualizar o texto para o próximo agente/task
                    texto += '\n\n ########' + str(resultado_task)

                if node['type'] == 'guardrail':
                    # Montar o prompt para o guardrail
                    prompt = f"""### GUARDRAIL ###
                    <tarefa>
                    {node['config']['rules']}
                    </tarefa>

                    ### TEXTO PARA PROCESSAR ###
                    <texto_processar>
                    {respostas[-1]['resposta']}
                    </texto_processar>

                    ### FORMATO DE SAÍDA ###
                    <formato_saida>
                    REESCREVA o <texto_processar> observando o formato de saída: {node['config']['outputFormat']}
                    
                    ATENÇÃO.
                    ESCREVA APENAS O TEXTO REESCRITO, SEM NADA MAIS. Se o formato de saída for JSON, escreva apenas o JSON, sem aspas.
                    Se o formato de saída for Texto, escreva apenas o texto, sem comentários nem nada além do texto.
                    </formato_saida>
                    """

                    # Executar o guardrail
                    resultado = IaController.groq(prompt)

                    # Extrair a resposta do resultado
                    if isinstance(resultado, tuple) and len(resultado) == 2:
                        if hasattr(resultado[0], 'get_json'):
                            resultado_data = resultado[0].get_json()
                        else:
                            resultado_data = resultado[0]
                    elif isinstance(resultado, dict):
                        resultado_data = resultado
                    elif hasattr(resultado, 'get_json'):
                        resultado_data = resultado.get_json()
                    else:
                        resultado_data = resultado

                    if isinstance(resultado_data, dict) and 'resposta' in resultado_data:
                        resposta_texto = resultado_data['resposta']
                    else:
                        resposta_texto = str(resultado_data)

                    # Criar registro de execução para este guardrail
                    execucao_guardrail = ExecucaoEquipe(
                        equipe_id=equipe_id,
                        sprint=execucao.sprint,
                        contexto=prompt,  # Salvar o prompt completo como contexto
                        resposta=resposta_texto
                    )
                    db.session.add(execucao_guardrail)
                    db.session.commit()

                    texto += f'<node_{node["id"]}>{resposta_texto}</node_{node["id"]}>\n\n\n'
                    respostas.append({
                        'resposta': resposta_texto,
                        'node': node
                    })

                # ########################################################################################
                # FORMATO DE SAÍDA
                # ########################################################################################

                if node['type'] == 'formatoSaida':
                    # Obter o conteúdo da última resposta
                    ultima_resposta = respostas[-1]['resposta']

                    # Se for uma tupla (Response, status_code), extrair o conteúdo
                    if isinstance(ultima_resposta, tuple):
                        if hasattr(ultima_resposta[0], 'get_data'):
                            conteudo_para_formatar = ultima_resposta[0].get_data(as_text=True)
                        else:
                            conteudo_para_formatar = str(ultima_resposta[0])
                    else:
                        conteudo_para_formatar = str(ultima_resposta)

                    # Gerar arquivo no formato especificado
                    formato = node['config']['type']
                    nome_arquivo = f"resultado_{equipe.nome}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

                    # Criar o contexto do formatoSaida
                    contexto_formato = f"""### FORMATO DE SAÍDA ###
                    <formato>
                    {formato}
                    </formato>

                    <conteudo_original>
                    {conteudo_para_formatar}
                    </conteudo_original>"""

                    arquivo_output, nome_arquivo_final, content_type = AgentesController.gerar_arquivo_saida(
                        conteudo_para_formatar, formato, nome_arquivo, equipe_id
                    )

                    # Criar registro de execução para este formatoSaida
                    execucao_formato = ExecucaoEquipe(
                        equipe_id=equipe_id,
                        sprint=execucao.sprint,
                        contexto=contexto_formato,
                        resposta=f"Arquivo gerado: {nome_arquivo_final} ({content_type})"
                    )
                    db.session.add(execucao_formato)
                    db.session.commit()

                    # Criar resposta para download
                    response = make_response(arquivo_output.getvalue())
                    response.headers['Content-Disposition'] = f'attachment; filename={nome_arquivo_final}'
                    response.headers['Content-Type'] = content_type

                    # Adicionar à lista de respostas
                    respostas.append({
                        'resposta': f"Arquivo gerado: {nome_arquivo_final}",
                        'node': node,
                        'arquivo_download': response
                    })

                    # Atualizar o texto para o próximo agente/task
                    texto += f'\n\n<node_{node["id"]}>Arquivo gerado: {nome_arquivo_final}</node_{node["id"]}>'


                # ########################################################################################
                # TEMPLATE
                # ########################################################################################
                if node['type'] == 'template':
                    ultima_resposta = respostas[-1]['resposta'] if respostas else ''

                    # Use no_autoflush para evitar flush automático
                    with db.session.no_autoflush:
                        template_arquivo = TemplateArquivo.query.filter_by(
                            node_id=node['id'], equipe_id=str(equipe.id)).first()

                    if template_arquivo:
                        # Criar o contexto do template
                        contexto_template = f"""### TEMPLATE ###
                        <template_arquivo>
                        {template_arquivo.filename}
                        </template_arquivo>

                        <conteudo_para_inserir>
                        {ultima_resposta}
                        </conteudo_para_inserir>"""

                        doc = Document(BytesIO(template_arquivo.data))
                        # Quebra as linhas em parágrafos
                        for linha in ultima_resposta.split('\n'):
                            doc.add_paragraph(linha)

                        output = BytesIO()
                        doc.save(output)
                        output.seek(0)
                        
                        nome_arquivo = f"resultado_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

                        # Criar registro de execução para este template
                        execucao_template = ExecucaoEquipe(
                            equipe_id=equipe_id,
                            sprint=execucao.sprint,
                            contexto=contexto_template,
                            resposta=f"Arquivo gerado: {nome_arquivo}"
                        )
                        db.session.add(execucao_template)
                        db.session.commit()
                        
                        # Criar resposta para download
                        response = make_response(output.getvalue())
                        response.headers['Content-Disposition'] = f'attachment; filename={nome_arquivo}'
                        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        return response

            # Criar resumo final da execução
            resumo_final = f"Execução completa. Total de etapas: {len(respostas)}. Última resposta: {respostas[-1]['resposta']}..."
            execucao.resposta = resumo_final
            db.session.commit()

            return render_template('agentes/resultado.html', resultado=respostas[-1]['resposta'], respostas=respostas, equipe=equipe)

    @staticmethod
    def listar_equipes_inteligentes():
        """Lista todas as equipes inteligentes salvas"""
        equipes = EquipeInteligente.query.order_by(
            EquipeInteligente.criado_em.desc()).all()
        
        for i, equipe in enumerate(equipes):
            layout = json.loads(equipe.layout) if isinstance(equipe.layout, str) else equipe.layout
            if layout['nodes'][-1]['type'] == 'formatoSaida':
                equipes[i].formato_saida = layout['nodes'][-1]['config']['type']
            elif layout['nodes'][-1]['type'] == 'template':
                equipes[i].formato_saida = 'Template do Word'
            else:
                equipes[i].formato_saida = 'Não definido'
        
        return render_template('agentes/equipes.html', equipes=equipes)

    @staticmethod
    def editar_equipe(id):
        """Página para editar uma equipe"""
        equipe = EquipeInteligente.query.get(id)

        return render_template('agentes/canva.html', equipe=equipe)

    @staticmethod
    def deletar_equipe(id):
        """Deleta uma equipe"""
        try:
            equipe = EquipeInteligente.query.get(id)
            db.session.delete(equipe)
            db.session.commit()
            return redirect(url_for('listar_equipes_inteligentes'))
        except Exception as e:
            db.session.rollback()
            print(f"Erro ao deletar equipe: {e}")
            flash(f'Erro ao deletar equipe: {str(e)}', 'danger')
            return jsonify({'status': 'error', 'message': f'Erro ao deletar equipe: {str(e)}'}), 500

    @staticmethod
    def download_equipe(id):
        """Download de uma equipe"""
        equipe = EquipeInteligente.query.get(id)

        # Prepara um dicionário com os dados da equipe
        equipe_data = {
            'nome': equipe.nome,
            'descricao': equipe.descricao,
            'processo': equipe.processo,
            'layout': json.loads(equipe.layout)
        }

        # Retorna um arquivo JSON para download
        response = make_response(json.dumps(equipe_data))
        response.headers['Content-Disposition'] = f'attachment; filename={equipe.nome.replace(" ", "_")}.json'
        response.headers['Content-Type'] = 'application/json'
        return response

    @staticmethod
    def importar_equipe():
        """Importa uma equipe"""
        arquivo = request.files['arquivo']

        # Ler o conteúdo do arquivo
        conteudo = arquivo.read()

        # Converter o conteúdo para JSON
        equipe_data = json.loads(conteudo)

        # Criar a equipe
        equipe = EquipeInteligente(
            nome=equipe_data['nome'],
            descricao=equipe_data['descricao'],
            processo=equipe_data['processo'],
            layout=json.dumps(equipe_data['layout'])
        )

        db.session.add(equipe)
        db.session.commit()

        return redirect(url_for('listar_equipes_inteligentes'))

    @staticmethod
    def gerar_arquivo_saida(conteudo, formato, nome_arquivo="resultado", equipe_id=None):
        """
        Gera um arquivo de saída no formato especificado

        Args:
            conteudo: Conteúdo a ser formatado
            formato: Tipo de formato ('texto', 'word', 'csv', 'pdf', 'mapamental', 'markdown', 'json')
            nome_arquivo: Nome base do arquivo
            equipe_id: ID da equipe (necessário para formato 'mapamental')

        Returns:
            BytesIO: Arquivo em memória pronto para download
        """
        try:
            # Garantir que o conteúdo seja uma string
            if not isinstance(conteudo, str):
                conteudo = str(conteudo)
            
            # Se o conteúdo estiver vazio, usar uma mensagem padrão
            if not conteudo or conteudo.strip() == "":
                conteudo = "Nenhum conteúdo disponível para gerar o arquivo."
            
            if formato == 'texto':
                # Arquivo de texto simples
                output = BytesIO()
                output.write(conteudo.encode('utf-8'))
                output.seek(0)
                return output, f"{nome_arquivo}.txt", "text/plain"

            elif formato == 'word':
                # Documento Word
                doc = Document()
                doc.add_paragraph(conteudo)

                output = BytesIO()
                doc.save(output)
                output.seek(0)
                return output, f"{nome_arquivo}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

            elif formato == 'csv':
                # Arquivo CSV
                output = BytesIO()
                writer = csv.writer(output)

                # Se o conteúdo for uma string, dividir por linhas
                if isinstance(conteudo, str):
                    linhas = conteudo.split('\n')
                    for linha in linhas:
                        if linha.strip():  # Ignorar linhas vazias
                            writer.writerow([linha])
                else:
                    # Se for uma lista ou outro formato
                    writer.writerow([conteudo])

                output.seek(0)
                return output, f"{nome_arquivo}.csv", "text/csv"

            elif formato == 'pdf':
                # Arquivo PDF (usando matplotlib para simplicidade)
                fig, ax = plt.subplots(figsize=(12, 8))
                ax.text(0.1, 0.9, 'Resultado do Processamento', transform=ax.transAxes,
                        fontsize=16, fontweight='bold')
                ax.text(0.1, 0.8, conteudo, transform=ax.transAxes, fontsize=10,
                        verticalalignment='top', wrap=True)
                ax.set_xlim(0, 1)
                ax.set_ylim(0, 1)
                ax.axis('off')

                output = BytesIO()
                plt.savefig(output, format='pdf', bbox_inches='tight')
                plt.close()
                output.seek(0)
                return output, f"{nome_arquivo}.pdf", "application/pdf"

            elif formato == 'mapamental':
                # Para formato mapamental, encontrar o último nó da equipe
                # e substituir {data} pelo conteúdo no campo texto
                
                # Obter a equipe atual para acessar o layout
                from models.models import EquipeInteligente
                equipe_atual = EquipeInteligente.query.get(equipe_id) if equipe_id else None
                
                if equipe_atual and equipe_atual.layout:
                    layout_equipe = json.loads(equipe_atual.layout) if isinstance(equipe_atual.layout, str) else equipe_atual.layout
                    
                    # Encontrar o último nó da equipe
                    if layout_equipe and 'nodes' in layout_equipe and layout_equipe['nodes']:
                        # Pegar o último nó da lista
                        ultimo_no = layout_equipe['nodes'][-1]
                        
                        # Verificar se o nó tem o campo 'config' e dentro dele 'description'
                        if 'config' in ultimo_no and 'description' in ultimo_no['config']:
                            texto_template = ultimo_no['config']['description']
                            
                            # Substituir {data} pelo conteúdo
                            if '{data}' in texto_template:
                                resultado_mapamental = texto_template.replace('{data}', conteudo)
                            else:
                                # Se não encontrar {data}, apenas concatenar o conteúdo
                                resultado_mapamental = texto_template + '\n\n' + conteudo
                        else:
                            # Se não houver campo 'description' em config, usar apenas o conteúdo
                            resultado_mapamental = conteudo
                    else:
                        # Se não houver nós, usar apenas o conteúdo
                        resultado_mapamental = conteudo
                else:
                    # Se não conseguir acessar a equipe, usar apenas o conteúdo
                    resultado_mapamental = conteudo
                
                # Retornar como HTML
                output = BytesIO()
                output.write(resultado_mapamental.encode('utf-8'))
                output.seek(0)
                return output, f"{nome_arquivo}.html", "text/html"

            elif formato == 'markdown':
                # Arquivo Markdown
                md_content = f"# Resultado do Processamento\n\n{conteudo}"
                output = BytesIO()
                output.write(md_content.encode('utf-8'))
                output.seek(0)
                return output, f"{nome_arquivo}.md", "text/markdown"

            elif formato == 'json':
                # Arquivo JSON
                try:
                    # Tentar fazer parse do conteúdo como JSON
                    if isinstance(conteudo, str):
                        json_data = json.loads(conteudo)
                    else:
                        json_data = conteudo
                except:
                    # Se não for JSON válido, criar um objeto com o conteúdo
                    json_data = {"resultado": conteudo,
                                 "timestamp": datetime.now().isoformat()}

                output = BytesIO()
                json.dump(json_data, output, indent=2, ensure_ascii=False)
                output.seek(0)
                return output, f"{nome_arquivo}.json", "application/json"

            else:
                # Formato não reconhecido, retornar como texto
                output = BytesIO()
                output.write(conteudo.encode('utf-8'))
                output.seek(0)
                return output, f"{nome_arquivo}.txt", "text/plain"

        except Exception as e:
            print(f"Erro ao gerar arquivo {formato}: {e}")
            # Em caso de erro, retornar como texto
            output = BytesIO()
            output.write(
                f"Erro ao gerar arquivo: {str(e)}\n\nConteúdo original:\n{conteudo}".encode('utf-8'))
            output.seek(0)
            return output, f"{nome_arquivo}_erro.txt", "text/plain"

    @staticmethod
    def download_arquivo_formato(equipe_id, node_id, formato):
        """Download de arquivo gerado por um nó formatoSaida"""
        try:
            equipe = EquipeInteligente.query.get_or_404(equipe_id)

            # Carregar o layout da equipe
            layout = json.loads(equipe.layout) if isinstance(
                equipe.layout, str) else equipe.layout

            # Encontrar o nó específico
            node = None
            for n in layout['nodes']:
                if n['id'] == node_id and n['type'] == 'formatoSaida':
                    node = n
                    break

            if not node:
                return jsonify({'error': 'Nó formatoSaida não encontrado'}), 404

            # Buscar a execução específica do formatoSaida
            from models.models import ExecucaoEquipe
            execucao_formato = ExecucaoEquipe.query.filter_by(
                equipe_id=equipe_id,
                contexto=f"Execução {equipe_id} - Agente {node_id} (formatoSaida)"
            ).order_by(ExecucaoEquipe.criado_em.desc()).first()
            
            if not execucao_formato:
                # Fallback: buscar a última execução da equipe que não seja resumo
                ultima_execucao = ExecucaoEquipe.query.filter_by(
                    equipe_id=equipe_id
                ).filter(
                    ~ExecucaoEquipe.resposta.like('%Execução completa%')
                ).order_by(ExecucaoEquipe.criado_em.desc()).first()
                
                if not ultima_execucao:
                    return jsonify({'error': 'Nenhuma execução encontrada para esta equipe'}), 404
                
                conteudo_real = ultima_execucao.resposta if ultima_execucao.resposta else "Nenhum conteúdo disponível"
            else:
                conteudo_real = execucao_formato.resposta if execucao_formato.resposta else "Nenhum conteúdo disponível"

            # Gerar o arquivo com o conteúdo real
            nome_arquivo = f"resultado_{equipe.nome}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            
            # Debug: imprimir informações sobre o conteúdo
            print(f"DEBUG: Gerando arquivo {formato} para equipe {equipe.nome}")
            print(f"DEBUG: Tamanho do conteúdo: {len(conteudo_real)} caracteres")
            print(f"DEBUG: Primeiros 100 caracteres: {conteudo_real[:100]}")
            
            arquivo_output, nome_arquivo_final, content_type = AgentesController.gerar_arquivo_saida(
                conteudo_real, formato, nome_arquivo, equipe_id
            )

            # Criar resposta para download
            response = make_response(arquivo_output.getvalue())
            response.headers['Content-Disposition'] = f'attachment; filename={nome_arquivo_final}'
            response.headers['Content-Type'] = content_type

            return response

        except Exception as e:
            print(f"Erro ao gerar download: {e}")
            return jsonify({'error': f'Erro ao gerar arquivo: {str(e)}'}), 500

    @staticmethod
    def agentes_menu(id):
        """Renderiza a página de agentes do menu dinâmico"""
        equipe = EquipeInteligente.query.get_or_404(id)

        # Fazer o parse do JSON do layout para que seja acessível no template
        layout = json.loads(equipe.layout) if isinstance(
            equipe.layout, str) else equipe.layout

        return render_template('agentes/executar_tarefas.html', equipe=equipe, trigger=True, layout=layout)

    @staticmethod
    def upload_template():
        print('--- [UPLOAD TEMPLATE] Iniciando upload_template ---')
        node_id = request.form.get('node_id')
        equipe_id = request.form.get('equipe_id')
        file = request.files.get('file')
        print(f'node_id: {node_id}, equipe_id: {equipe_id}, file: {file}')
        if not file:
            print('[UPLOAD TEMPLATE] Nenhum arquivo enviado!')
            return jsonify({'status': 'error', 'message': 'Arquivo não enviado'}), 400

        try:
            template = TemplateArquivo(
                node_id=node_id,
                equipe_id=equipe_id,
                filename=file.filename,
                data=file.read(),
                mimetype=file.mimetype
            )
            db.session.add(template)
            db.session.commit()
            print(
                f'[UPLOAD TEMPLATE] Template salvo com sucesso! node_id={node_id}, equipe_id={equipe_id}, filename={file.filename}')
            return jsonify({'status': 'success'})
        except Exception as e:
            print(f'[UPLOAD TEMPLATE] Erro ao salvar template: {e}')
            return jsonify({'status': 'error', 'message': str(e)})

    @staticmethod
    def listar_execucoes(equipe_id=None):
        """Lista as execuções de uma equipe específica ou todas as execuções com paginação e filtros"""
        try:
            # Parâmetros de paginação
            page = request.args.get('page', 1, type=int)
            per_page = request.args.get('per_page', 10, type=int)

            # Parâmetros de filtro
            filtro_equipe = request.args.get('equipe', type=int)
            filtro_sprint = request.args.get('sprint', type=int)
            filtro_contexto = request.args.get('contexto', '')
            filtro_resposta = request.args.get('resposta', '')

            # Construir a query base
            query = ExecucaoEquipe.query

            # Aplicar filtros
            if equipe_id:
                query = query.filter_by(equipe_id=equipe_id)
            elif filtro_equipe:
                query = query.filter_by(equipe_id=filtro_equipe)

            if filtro_sprint:
                query = query.filter_by(sprint=filtro_sprint)

            if filtro_contexto:
                query = query.filter(ExecucaoEquipe.contexto.like(f'%{filtro_contexto}%'))

            if filtro_resposta:
                query = query.filter(ExecucaoEquipe.resposta.like(f'%{filtro_resposta}%'))

            # Ordenar e paginar
            execucoes_paginadas = query.order_by(ExecucaoEquipe.criado_em.desc()).paginate(
                page=page, 
                per_page=per_page,
                error_out=False
            )

            # Buscar todas as equipes para o filtro
            equipes = EquipeInteligente.query.all()

            # Buscar sprints únicos para o filtro
            sprints = db.session.query(ExecucaoEquipe.sprint).distinct().order_by(ExecucaoEquipe.sprint).all()
            sprints = [s[0] for s in sprints if s[0] is not None]  # Remover None e extrair valores

            # Se estiver filtrando por equipe específica
            equipe_atual = None
            if equipe_id:
                equipe_atual = EquipeInteligente.query.get_or_404(equipe_id)

            return render_template(
                'agentes/execucoes.html',
                execucoes=execucoes_paginadas,
                equipe=equipe_atual,
                equipes=equipes,
                sprints=sprints,
                filtros={
                    'equipe': filtro_equipe,
                    'sprint': filtro_sprint,
                    'contexto': filtro_contexto,
                    'resposta': filtro_resposta
                }
            )

        except Exception as e:
            print(f"ERRO em listar_execucoes: {e}")
            return render_template('agentes/execucoes.html', execucoes=[], error=str(e))

    @staticmethod
    def detalhes_execucao(execucao_id):
        """Mostra os detalhes de uma execução específica e todas as execuções da mesma sprint"""
        execucao = ExecucaoEquipe.query.get_or_404(execucao_id)
        
        # Buscar todas as execuções da mesma sprint
        execucoes_sprint = ExecucaoEquipe.query.filter_by(
            equipe_id=execucao.equipe_id,
            sprint=execucao.sprint
        ).order_by(ExecucaoEquipe.criado_em.asc()).all()
        
        return render_template('agentes/detalhes_execucao.html', 
                             execucao=execucao, 
                             execucoes_sprint=execucoes_sprint)

    @staticmethod
    def excluir_execucao(execucao_id):
        """Exclui uma execução específica"""
        try:
            execucao = ExecucaoEquipe.query.get_or_404(execucao_id)
            equipe_id = execucao.equipe_id  # Salvar o ID da equipe antes de excluir
            db.session.delete(execucao)
            db.session.commit()
            
            # Redirecionar para a lista de execuções da equipe
            return redirect(f'/agentesinteligentes/equipe/{equipe_id}/execucoes')
        except Exception as e:
            db.session.rollback()
            print(f"Erro ao excluir execução: {e}")
            flash(f'Erro ao excluir execução: {str(e)}', 'error')
            return redirect('/agentesinteligentes/execucoes')

    @staticmethod
    def exportar_execucoes(equipe_id=None):
        """Exporta execuções para CSV"""
        try:
            if equipe_id:
                execucoes = ExecucaoEquipe.query.filter_by(equipe_id=equipe_id).order_by(ExecucaoEquipe.criado_em.desc()).all()
                equipe = EquipeInteligente.query.get_or_404(equipe_id)
                filename = f"execucoes_equipe_{equipe.nome}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"
            else:
                execucoes = ExecucaoEquipe.query.order_by(ExecucaoEquipe.criado_em.desc()).all()
                filename = f"todas_execucoes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv"

            # Criar arquivo CSV
            output = BytesIO()
            writer = csv.writer(output)
            
            # Cabeçalho
            writer.writerow(['ID', 'Equipe', 'Contexto', 'Resposta', 'Data Criação', 'Última Atualização'])
            
            # Dados
            for execucao in execucoes:
                writer.writerow([
                    execucao.id,
                    execucao.equipe.nome,
                    execucao.contexto[:500] + '...' if len(execucao.contexto) > 500 else execucao.contexto,
                    execucao.resposta[:500] + '...' if len(execucao.resposta) > 500 else execucao.resposta,
                    execucao.criado_em.strftime('%d/%m/%Y %H:%M:%S'),
                    execucao.atualizado_em.strftime('%d/%m/%Y %H:%M:%S')
                ])
            
            output.seek(0)
            
            # Criar resposta para download
            response = make_response(output.getvalue())
            response.headers['Content-Disposition'] = f'attachment; filename={filename}'
            response.headers['Content-Type'] = 'text/csv'
            
            return response
            
        except Exception as e:
            print(f"Erro ao exportar execuções: {e}")
            return jsonify({'status': 'error', 'message': f'Erro ao exportar execuções: {str(e)}'}), 500
