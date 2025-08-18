from flask import render_template, request, jsonify, send_file
import tempfile
import os
import requests
import PyPDF2
from PyPDF2 import PdfReader, PdfWriter
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
import io
import json
from datetime import datetime
import controllers.IaController as IaController
import zipfile


class PdfController:
    @staticmethod
    def index():
        """Página inicial das ferramentas de PDF"""
        return render_template('pdf/index.html')

    @staticmethod
    def unir_pdfs():
        """Unir vários PDFs em um só"""
        if request.method == 'GET':
            return render_template('pdf/unir.html')
        
        try:
            files = request.files.getlist('pdfs')
            if not files or all(file.filename == '' for file in files):
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Criar um PDF writer
            pdf_writer = PdfWriter()
            
            # Processar cada arquivo PDF
            for file in files:
                if file.filename and file.filename.lower().endswith('.pdf'):
                    pdf_reader = PdfReader(file)
                    for page in pdf_reader.pages:
                        pdf_writer.add_page(page)

            # Salvar o PDF unido
            output = io.BytesIO()
            pdf_writer.write(output)
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'pdf_unido_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                mimetype='application/pdf'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao unir PDFs: {str(e)}'
            }), 500

    @staticmethod
    def separar_pdf():
        """Separar um PDF em vários por página, tamanho, página específica ou quantidade de páginas"""
        if request.method == 'GET':
            return render_template('pdf/separar.html')
        
        try:
            file = request.files['pdf']
            tipo_separacao = request.form.get('tipo_separacao')
            
            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            pdf_reader = PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            # Criar um buffer para o ZIP
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                
                if tipo_separacao == 'por_pagina':
                    # Separar por página individual
                    for page_num in range(total_pages):
                        pdf_writer = PdfWriter()
                        pdf_writer.add_page(pdf_reader.pages[page_num])
                        
                        # Criar PDF em memória
                        pdf_buffer = io.BytesIO()
                        pdf_writer.write(pdf_buffer)
                        pdf_buffer.seek(0)
                        
                        # Adicionar ao ZIP
                        zip_file.writestr(f'pagina_{page_num + 1}.pdf', pdf_buffer.getvalue())
                        
                elif tipo_separacao == 'por_tamanho':
                    # Separar por tamanho máximo
                    tamanho_max = int(request.form.get('tamanho_max', 5)) * 1024 * 1024  # Converter MB para bytes
                    current_pdf = PdfWriter()
                    current_size = 0
                    file_counter = 1
                    
                    for page_num in range(total_pages):
                        page = pdf_reader.pages[page_num]
                        
                        # Estimar tamanho da página (aproximação)
                        page_buffer = io.BytesIO()
                        temp_writer = PdfWriter()
                        temp_writer.add_page(page)
                        temp_writer.write(page_buffer)
                        page_size = len(page_buffer.getvalue())
                        
                        if current_size + page_size > tamanho_max and current_size > 0:
                            # Salvar PDF atual e começar novo
                            pdf_buffer = io.BytesIO()
                            current_pdf.write(pdf_buffer)
                            pdf_buffer.seek(0)
                            zip_file.writestr(f'parte_{file_counter}.pdf', pdf_buffer.getvalue())
                            
                            # Resetar para próxima parte
                            current_pdf = PdfWriter()
                            current_size = 0
                            file_counter += 1
                        
                        current_pdf.add_page(page)
                        current_size += page_size
                    
                    # Salvar última parte se houver
                    if current_size > 0:
                        pdf_buffer = io.BytesIO()
                        current_pdf.write(pdf_buffer)
                        pdf_buffer.seek(0)
                        zip_file.writestr(f'parte_{file_counter}.pdf', pdf_buffer.getvalue())
                        
                elif tipo_separacao == 'pagina_especifica':
                    # Separar em uma página específica
                    pagina_separacao = int(request.form.get('pagina_separacao', 1)) - 1  # Converter para índice baseado em 0
                    
                    if pagina_separacao >= total_pages:
                        return jsonify({
                            'status': 'error',
                            'message': f'Página especificada ({pagina_separacao + 1}) é maior que o total de páginas ({total_pages})'
                        }), 400
                    
                    # Primeira parte (até a página especificada)
                    pdf_writer1 = PdfWriter()
                    for i in range(pagina_separacao + 1):
                        pdf_writer1.add_page(pdf_reader.pages[i])
                    
                    pdf_buffer1 = io.BytesIO()
                    pdf_writer1.write(pdf_buffer1)
                    pdf_buffer1.seek(0)
                    zip_file.writestr('parte_1.pdf', pdf_buffer1.getvalue())
                    
                    # Segunda parte (após a página especificada)
                    if pagina_separacao + 1 < total_pages:
                        pdf_writer2 = PdfWriter()
                        for i in range(pagina_separacao + 1, total_pages):
                            pdf_writer2.add_page(pdf_reader.pages[i])
                        
                        pdf_buffer2 = io.BytesIO()
                        pdf_writer2.write(pdf_buffer2)
                        pdf_buffer2.seek(0)
                        zip_file.writestr('parte_2.pdf', pdf_buffer2.getvalue())
                        
                elif tipo_separacao == 'quantidade_paginas':
                    # Separar por quantidade de páginas por arquivo
                    qtd_paginas = int(request.form.get('qtd_paginas', 1))
                    
                    if qtd_paginas <= 0:
                        return jsonify({
                            'status': 'error',
                            'message': 'Quantidade de páginas deve ser maior que zero'
                        }), 400
                    
                    file_counter = 1
                    current_pdf = PdfWriter()
                    pages_in_current = 0
                    
                    for page_num in range(total_pages):
                        current_pdf.add_page(pdf_reader.pages[page_num])
                        pages_in_current += 1
                        
                        if pages_in_current >= qtd_paginas:
                            # Salvar PDF atual
                            pdf_buffer = io.BytesIO()
                            current_pdf.write(pdf_buffer)
                            pdf_buffer.seek(0)
                            zip_file.writestr(f'parte_{file_counter}.pdf', pdf_buffer.getvalue())
                            
                            # Resetar para próxima parte
                            current_pdf = PdfWriter()
                            pages_in_current = 0
                            file_counter += 1
                    
                    # Salvar última parte se houver páginas restantes
                    if pages_in_current > 0:
                        pdf_buffer = io.BytesIO()
                        current_pdf.write(pdf_buffer)
                        pdf_buffer.seek(0)
                        zip_file.writestr(f'parte_{file_counter}.pdf', pdf_buffer.getvalue())

            # Retornar o ZIP
            zip_buffer.seek(0)
            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name=f'pdf_separado_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip',
                mimetype='application/zip'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao separar PDF: {str(e)}'
            }), 500

    @staticmethod
    def girar_pdf():
        """Girar todas as páginas ou uma página específica"""
        if request.method == 'GET':
            return render_template('pdf/girar.html')
        
        try:
            file = request.files['pdf']
            angulo = int(request.form.get('angulo', 90))
            tipo_rotacao = request.form.get('tipo_rotacao', 'todas')
            pagina_especifica = int(request.form.get('pagina_especifica', 1)) - 1

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            pdf_reader = PdfReader(file)
            pdf_writer = PdfWriter()

            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                
                if tipo_rotacao == 'todas':
                    page.rotate(angulo)
                elif tipo_rotacao == 'especifica' and page_num == pagina_especifica:
                    page.rotate(angulo)
                
                pdf_writer.add_page(page)

            output = io.BytesIO()
            pdf_writer.write(output)
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'pdf_rotacionado_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                mimetype='application/pdf'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao girar PDF: {str(e)}'
            }), 500

    @staticmethod
    def dividir_por_bookmarks():
        """Dividir PDF por bookmarks"""
        if request.method == 'GET':
            return render_template('pdf/bookmarks.html')
        
        try:
            file = request.files['pdf']
            nivel_bookmark = int(request.form.get('nivel_bookmark', 1))

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Usar PyMuPDF para acessar bookmarks
            doc = fitz.open(stream=file.read(), filetype="pdf")
            toc = doc.get_toc()
            
            # Implementar lógica de divisão por bookmarks
            # Por enquanto, retornar apenas o PDF original
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'pdf_bookmarks_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                mimetype='application/pdf'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao dividir por bookmarks: {str(e)}'
            }), 500

    @staticmethod
    def extrair_texto():
        """Extrair texto do PDF"""
        if request.method == 'GET':
            return render_template('pdf/extrair_texto.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            pdf_reader = PdfReader(file)
            texto_completo = ""

            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            return jsonify({
                'status': 'success',
                'texto': texto_completo
            })

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao extrair texto: {str(e)}'
            }), 500

    @staticmethod
    def resumir_pdf():
        """Resumir PDF com IA"""
        if request.method == 'GET':
            return render_template('pdf/resumir.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Usar IA para resumir
            prompt = f"Por favor, faça um resumo completo e bem estruturado do seguinte texto: {texto_completo[:4000]}"
            resultado = IaController.groq(prompt, None, None)
            
            # Extrair a resposta do resultado do IaController
            if isinstance(resultado, tuple) and len(resultado) == 2:
                # Se for uma tupla (Response, status_code), extrair o conteúdo
                if hasattr(resultado[0], 'get_json'):
                    resultado_data = resultado[0].get_json()
                else:
                    resultado_data = resultado[0]
            elif isinstance(resultado, dict):
                resultado_data = resultado
            elif hasattr(resultado, 'get_json'):
                resultado_data = resultado.get_json()
            else:
                resultado_data = resultado

            # Verificar se foi bem-sucedido
            if isinstance(resultado_data, dict) and resultado_data.get('status') == 'success':
                return jsonify({
                    'status': 'success',
                    'response': resultado_data.get('resposta', 'Resumo gerado com sucesso!')
                })
            else:
                return jsonify({
                    'status': 'error',
                    'message': resultado_data.get('message', 'Erro ao gerar resumo')
                }), 500

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao resumir PDF: {str(e)}'
            }), 500

    @staticmethod
    def pdf_para_markdown():
        """Transformar PDF em Markdown com IA"""
        if request.method == 'GET':
            return render_template('pdf/markdown.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Usar IA para converter em markdown
            prompt = f"Converta o seguinte texto em um documento markdown bem estruturado, mantendo a hierarquia e formatação adequada: {texto_completo[:4000]}"
            resultado = IaController.groq(prompt, None, None)

            # Extrair a resposta do resultado do IaController
            if isinstance(resultado, tuple) and len(resultado) == 2:
                # Se for uma tupla (Response, status_code), extrair o conteúdo
                if hasattr(resultado[0], 'get_json'):
                    resultado_data = resultado[0].get_json()
                else:
                    resultado_data = resultado[0]
            elif isinstance(resultado, dict):
                resultado_data = resultado
            elif hasattr(resultado, 'get_json'):
                resultado_data = resultado.get_json()
            else:
                resultado_data = resultado

            # Verificar se foi bem-sucedido
            if isinstance(resultado_data, dict) and resultado_data.get('status') == 'success':
                return jsonify({
                    'status': 'success',
                    'response': resultado_data.get('resposta', 'Markdown gerado com sucesso!')
                })
            else:
                return jsonify({
                    'status': 'error',
                    'message': resultado_data.get('message', 'Erro ao converter para markdown')
                }), 500

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao converter para markdown: {str(e)}'
            }), 500

    @staticmethod
    def pdf_para_mapa_mental():
        """Transformar PDF em mapa mental com IA"""
        if request.method == 'GET':
            return render_template('pdf/mapa_mental.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Usar IA para criar mapa mental
            prompt = f"Crie um mapa mental estruturado do seguinte conteúdo, organizando as ideias principais e secundárias de forma hierárquica: {texto_completo[:4000]}"
            resultado = IaController.groq(prompt, None, None)

            # Extrair a resposta do resultado do IaController
            if isinstance(resultado, tuple) and len(resultado) == 2:
                # Se for uma tupla (Response, status_code), extrair o conteúdo
                if hasattr(resultado[0], 'get_json'):
                    resultado_data = resultado[0].get_json()
                else:
                    resultado_data = resultado[0]
            elif isinstance(resultado, dict):
                resultado_data = resultado
            elif hasattr(resultado, 'get_json'):
                resultado_data = resultado.get_json()
            else:
                resultado_data = resultado

            # Verificar se foi bem-sucedido
            if isinstance(resultado_data, dict) and resultado_data.get('status') == 'success':
                return jsonify({
                    'status': 'success',
                    'response': resultado_data.get('resposta', 'Mapa mental gerado com sucesso!')
                })
            else:
                return jsonify({
                    'status': 'error',
                    'message': resultado_data.get('message', 'Erro ao criar mapa mental')
                }), 500

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao criar mapa mental: {str(e)}'
            }), 500

    @staticmethod
    def pdf_para_docx():
        """Transformar PDF em DOCX"""
        if request.method == 'GET':
            return render_template('pdf/docx.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Criar documento DOCX
            doc = Document()
            doc.add_heading('Documento convertido de PDF', 0)
            doc.add_paragraph(texto_completo)

            # Salvar em buffer
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'documento_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao converter para DOCX: {str(e)}'
            }), 500

    @staticmethod
    def pdf_para_xlsx():
        """Transformar PDF em XLSX"""
        if request.method == 'GET':
            return render_template('pdf/xlsx.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Criar DataFrame e salvar como XLSX
            # Por simplicidade, vou criar uma planilha com o texto em uma coluna
            df = pd.DataFrame({'Conteúdo': [texto_completo]})
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='PDF_Convertido')
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'planilha_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx',
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao converter para XLSX: {str(e)}'
            }), 500

    @staticmethod
    def pdf_para_csv():
        """Transformar PDF em CSV"""
        if request.method == 'GET':
            return render_template('pdf/csv.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Extrair texto do PDF
            pdf_reader = PdfReader(file)
            texto_completo = ""
            for page in pdf_reader.pages:
                texto_completo += page.extract_text() + "\n"

            # Criar DataFrame e salvar como CSV
            df = pd.DataFrame({'Conteúdo': [texto_completo]})
            
            output = io.BytesIO()
            df.to_csv(output, index=False, encoding='utf-8')
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'dados_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv',
                mimetype='text/csv'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao converter para CSV: {str(e)}'
            }), 500 

    @staticmethod
    def inserir_pdf():
        """Inserir um PDF no meio de outro PDF"""
        if request.method == 'GET':
            return render_template('pdf/inserir.html')
        
        try:
            pdf_principal = request.files['pdf_principal']
            pdf_inserir = request.files['pdf_inserir']
            posicao = request.form.get('posicao', 'final')
            pagina_especifica = request.form.get('pagina_especifica', '1')

            if not pdf_principal or pdf_principal.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF principal foi enviado'
                }), 400

            if not pdf_inserir or pdf_inserir.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF para inserir foi enviado'
                }), 400

            # Ler o PDF principal
            pdf_reader_principal = PdfReader(pdf_principal)
            pdf_writer = PdfWriter()

            # Ler o PDF a ser inserido
            pdf_reader_inserir = PdfReader(pdf_inserir)

            # Processar baseado na posição escolhida
            if posicao == 'inicio':
                # Inserir no início
                for page in pdf_reader_inserir.pages:
                    pdf_writer.add_page(page)
                for page in pdf_reader_principal.pages:
                    pdf_writer.add_page(page)
                    
            elif posicao == 'pagina_especifica':
                # Inserir após uma página específica
                pagina = int(pagina_especifica) - 1  # Converter para índice baseado em 0
                
                # Adicionar páginas do PDF principal até a página especificada
                for i in range(min(pagina, len(pdf_reader_principal.pages))):
                    pdf_writer.add_page(pdf_reader_principal.pages[i])
                
                # Inserir páginas do PDF a ser inserido
                for page in pdf_reader_inserir.pages:
                    pdf_writer.add_page(page)
                
                # Adicionar as páginas restantes do PDF principal
                for i in range(pagina, len(pdf_reader_principal.pages)):
                    pdf_writer.add_page(pdf_reader_principal.pages[i])
                    
            else:  # final
                # Inserir no final
                for page in pdf_reader_principal.pages:
                    pdf_writer.add_page(page)
                for page in pdf_reader_inserir.pages:
                    pdf_writer.add_page(page)

            # Criar o arquivo de saída
            output = io.BytesIO()
            pdf_writer.write(output)
            output.seek(0)

            return send_file(
                output,
                as_attachment=True,
                download_name=f'pdf_inserido_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                mimetype='application/pdf'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao inserir PDF: {str(e)}'
            }), 500 

    @staticmethod
    def extrair_paginas():
        """Extrair páginas específicas de um PDF"""
        if request.method == 'GET':
            return render_template('pdf/extrair_paginas.html')
        
        try:
            file = request.files['pdf']
            paginas_str = request.form.get('paginas', '').strip()

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            if not paginas_str:
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhuma página foi especificada'
                }), 400

            # Ler o PDF
            pdf_reader = PdfReader(file)
            total_pages = len(pdf_reader.pages)

            # Processar a string de páginas
            # Substituir ; por ,
            paginas_str = paginas_str.replace(';', ',')
            
            # Fazer split em , e processar cada bloco
            blocos = [bloco.strip() for bloco in paginas_str.split(',') if bloco.strip()]
            
            # Conjunto para armazenar páginas únicas (evitar duplicatas)
            paginas_para_extrair = set()
            
            for bloco in blocos:
                if '-' in bloco:
                    # Se tem - entre os números, extrair desde a primeira até a segunda
                    try:
                        inicio, fim = map(int, bloco.split('-'))
                        # Validar se as páginas estão dentro do range
                        if inicio < 1 or fim > total_pages:
                            return jsonify({
                                'status': 'error',
                                'message': f'Páginas {inicio}-{fim} estão fora do range válido (1-{total_pages})'
                            }), 400
                        if inicio > fim:
                            return jsonify({
                                'status': 'error',
                                'message': f'Página inicial ({inicio}) não pode ser maior que a final ({fim})'
                            }), 400
                        
                        # Adicionar todas as páginas do intervalo (convertendo para índice baseado em 0)
                        for pagina in range(inicio, fim + 1):
                            paginas_para_extrair.add(pagina - 1)
                    except ValueError:
                        return jsonify({
                            'status': 'error',
                            'message': f'Formato inválido no intervalo: {bloco}'
                        }), 400
                else:
                    # Se for um número, extrair aquela página
                    try:
                        pagina = int(bloco)
                        if pagina < 1 or pagina > total_pages:
                            return jsonify({
                                'status': 'error',
                                'message': f'Página {pagina} está fora do range válido (1-{total_pages})'
                            }), 400
                        paginas_para_extrair.add(pagina - 1)  # Converter para índice baseado em 0
                    except ValueError:
                        return jsonify({
                            'status': 'error',
                            'message': f'Página inválida: {bloco}'
                        }), 400

            if not paginas_para_extrair:
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhuma página válida foi especificada'
                }), 400

            # Criar um buffer para o ZIP
            zip_buffer = io.BytesIO()
            
            with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # Ordenar as páginas para manter a ordem
                paginas_ordenadas = sorted(paginas_para_extrair)
                
                for i, pagina_idx in enumerate(paginas_ordenadas):
                    # Criar um PDF com apenas esta página
                    pdf_writer = PdfWriter()
                    pdf_writer.add_page(pdf_reader.pages[pagina_idx])
                    
                    # Criar PDF em memória
                    pdf_buffer = io.BytesIO()
                    pdf_writer.write(pdf_buffer)
                    pdf_buffer.seek(0)
                    
                    # Adicionar ao ZIP com nome descritivo
                    nome_arquivo = f'pagina_{pagina_idx + 1}.pdf'
                    zip_file.writestr(nome_arquivo, pdf_buffer.getvalue())

                # Agrupe todas as paginas em um unico PDF
                pdf_writer = PdfWriter()
                for pagina_idx in paginas_ordenadas:
                    pdf_writer.add_page(pdf_reader.pages[pagina_idx])
                    
                # Adicionar o PDF final ao ZIP
                pdf_buffer = io.BytesIO()
                pdf_writer.write(pdf_buffer)
                pdf_buffer.seek(0)
                zip_file.writestr('todas_paginas.pdf', pdf_buffer.getvalue())
            

            # Retornar o ZIP
            zip_buffer.seek(0)
            return send_file(
                zip_buffer,
                as_attachment=True,
                download_name=f'paginas_extraidas_{datetime.now().strftime("%Y%m%d_%H%M%S")}.zip',
                mimetype='application/zip'
            )

        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao extrair páginas: {str(e)}'
            }), 500

    @staticmethod
    def ocr_pdf():
        """Fazer OCR em PDF usando OCR.Space API"""
        if request.method == 'GET':
            return render_template('pdf/ocr.html')
        
        try:
            file = request.files['pdf']

            if not file or file.filename == '':
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF foi enviado'
                }), 400

            # Criar um arquivo temporário para salvar o PDF
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_file_path = temp_file.name
            temp_file.close()
            
            try:
                # Abrir e enviar o arquivo para a API
                with open(temp_file_path, 'rb') as f:
                    files = {
                        'file': (file.filename, f, 'application/pdf')
                    }
                    payload = {
                        'apikey': os.getenv('OCRSPACE_API_KEY'),
                        'language': 'por',
                        'isOverlayRequired': False,
                        'filetype': 'PDF',
                        'detectOrientation': True,
                        'scale': True,
                        'OCREngine': 2  # Engine mais preciso
                    }
                    
                    r = requests.post(
                        'https://api.ocr.space/parse/image',
                        files=files,
                        data=payload,
                    )
                    
                    # Verificar se a resposta é bem-sucedida
                    if r.status_code != 200:
                        raise Exception(f"Erro na API OCR.Space: Status {r.status_code}")
                    
                    try:
                        resultado_json = r.json()
                        if resultado_json.get('OCRExitCode') != 1:
                            erro = resultado_json.get('ErrorMessage', 'Erro desconhecido no OCR')
                            raise Exception(f"Erro no processamento OCR: {erro}")
                        
                        # Extrair o texto do resultado
                        texto_extraido = ""
                        for page_result in resultado_json.get('ParsedResults', []):
                            parsed_text = page_result.get('ParsedText', '')
                            if parsed_text:
                                texto_extraido += parsed_text + "\n\n"  # Adiciona espaço extra entre páginas
                        
                        if not texto_extraido.strip():
                            raise Exception("Nenhum texto foi extraído do documento")
                        
                        return jsonify({
                            'status': 'success',
                            'response': texto_extraido,
                            'message': 'OCR processado com sucesso!'
                        })
                        
                    except json.JSONDecodeError:
                        raise Exception("Resposta inválida da API OCR.Space")
                    
            finally:
                # Garantir que o arquivo temporário seja removido mesmo se houver erro
                try:
                    if os.path.exists(temp_file_path):
                        os.unlink(temp_file_path)
                except Exception as e:
                    print(f"Erro ao remover arquivo temporário: {e}")
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao fazer OCR: {str(e)}'
            }), 500


    @staticmethod
    def comprimir_pdf():
        """Comprimir PDF reduzindo qualidade de imagens e otimizando o arquivo"""
        if request.method == 'GET':
            return render_template('pdf/comprimir.html')
        
        # Configurações de compressão por nível
        compression_config = {
            'baixo': {'image_quality': 80, 'scale_factor': 0.9},
            'medio': {'image_quality': 60, 'scale_factor': 0.75},
            'alto': {'image_quality': 40, 'scale_factor': 0.6}
        }
        
        temp_file_path = None
        doc = None
        
        try:
            # Validar arquivo
            file = request.files.get('pdf')
            if not file or not file.filename or not file.filename.lower().endswith('.pdf'):
                return jsonify({
                    'status': 'error',
                    'message': 'Nenhum arquivo PDF válido foi enviado'
                }), 400

            # Obter configuração de compressão
            nivel_compressao = request.form.get('nivel_compressao', 'medio')
            if nivel_compressao not in compression_config:
                nivel_compressao = 'medio'
            
            config = compression_config[nivel_compressao]
            image_quality = config['image_quality']
            scale_factor = config['scale_factor']

            # Criar arquivo temporário
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            temp_file_path = temp_file.name
            file.save(temp_file_path)
            temp_file.close()
            
            # Obter tamanho original
            tamanho_original = os.path.getsize(temp_file_path)
            
            # Abrir e validar PDF
            doc = fitz.open(temp_file_path)
            if doc.page_count == 0:
                return jsonify({
                    'status': 'error',
                    'message': 'PDF inválido ou corrompido'
                }), 400

            # Comprimir imagens em todas as páginas
            for page_num in range(doc.page_count):
                page = doc[page_num]
                image_list = page.get_images(full=True)
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    try:
                        # Extrair imagem
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        pix = fitz.Pixmap(image_bytes)
                        
                        # Verificar se pode ser comprimida (sem transparência complexa)
                        if pix.n - pix.alpha < 4:
                            # Redimensionar se necessário
                            if scale_factor < 1.0:
                                mat = fitz.Matrix(scale_factor, scale_factor)
                                pix_resized = fitz.Pixmap(pix, mat)
                                pix.clear()
                                pix = pix_resized
                            
                            # Comprimir baseado no tipo da imagem
                            if pix.colorspace and pix.colorspace.name != "GRAY":
                                compressed_bytes = pix.tobytes("jpeg", image_quality)
                            else:
                                # Para escala de cinza, PNG pode ser melhor
                                compressed_bytes = pix.tobytes("png")
                            
                            # Atualizar imagem no documento
                            doc.update_stream(xref, compressed_bytes)
                        
                        # Limpar memória
                        pix.clear()
                        pix = None
                        
                    except Exception as e:
                        print(f"Aviso: Erro ao comprimir imagem {img_index} na página {page_num}: {e}")
                        continue

            # Otimizar conteúdo das páginas
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    page.clean_contents()
                except Exception as e:
                    print(f"Aviso: Erro ao otimizar página {page_num}: {e}")
                    continue

            # Gerar PDF otimizado
            output = io.BytesIO()
            doc.save(
                output,
                garbage=4,              # Nível máximo de limpeza
                deflate=True,           # Compressão deflate
                clean=True,             # Limpar conteúdo
                ascii=False,            # Manter encoding binário
                expand=0,               # Não expandir objetos
                linear=False,           # Não linearizar
                pretty=False,           # Não formatar para legibilidade
                deflate_images=True,    # Comprimir imagens
                deflate_fonts=True      # Comprimir fontes
            )
            output.seek(0)

            # Calcular estatísticas
            tamanho_comprimido = len(output.getvalue())
            reducao_percentual = round(((tamanho_original - tamanho_comprimido) / tamanho_original) * 100, 1) if tamanho_original > 0 else 0
            
            # Log das estatísticas
            print(f"=== Compressão PDF ===")
            print(f"Original: {tamanho_original:,} bytes ({tamanho_original/1024/1024:.2f} MB)")
            print(f"Comprimido: {tamanho_comprimido:,} bytes ({tamanho_comprimido/1024/1024:.2f} MB)")
            print(f"Redução: {reducao_percentual}%")
            print(f"Nível: {nivel_compressao} (Qualidade: {image_quality}%, Escala: {scale_factor})")
            print(f"======================")

            return send_file(
                output,
                as_attachment=True,
                download_name=f'pdf_comprimido_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf',
                mimetype='application/pdf'
            )

        except fitz.fitz.FileDataError:
            return jsonify({
                'status': 'error',
                'message': 'Arquivo PDF corrompido ou inválido'
            }), 500
        except MemoryError:
            return jsonify({
                'status': 'error',
                'message': 'Arquivo muito grande para processamento'
            }), 500
        except Exception as e:
            return jsonify({
                'status': 'error',
                'message': f'Erro ao comprimir PDF: {str(e)}'
            }), 500
        finally:
            # Cleanup garantido
            if doc:
                try:
                    doc.close()
                except Exception:
                    pass
            
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass